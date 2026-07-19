import docx
import fitz
import pandas as pd
import pytest

import parsers


def test_text_parser(tmp_path):
    f = tmp_path / "a.txt"
    f.write_text("hello world\nsecond line")
    result = parsers.extract(f)
    assert "hello world" in result.text
    assert result.needs_ocr is False
    assert result.page_count == 1


def test_docx_parser(tmp_path):
    f = tmp_path / "b.docx"
    d = docx.Document()
    d.add_paragraph("Docx paragraph one")
    d.add_paragraph("Docx paragraph two")
    d.save(str(f))
    result = parsers.extract(f)
    assert "Docx paragraph one" in result.text
    assert "Docx paragraph two" in result.text
    assert result.needs_ocr is False
    assert result.metadata["paragraph_count"] == 2


def test_pdf_parser_detects_normal_text_page(tmp_path):
    f = tmp_path / "c.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "This is a normal text PDF with plenty of extractable text. " * 3)
    doc.save(str(f))
    doc.close()

    result = parsers.extract(f, low_text_chars_per_page=50)
    assert result.needs_ocr is False
    assert result.page_count == 1
    assert len(result.text) > 50


def test_pdf_parser_flags_scanned_page_for_ocr(tmp_path):
    f = tmp_path / "d.pdf"
    doc = fitz.open()
    doc.new_page()  # blank page, no text layer — simulates a scanned page
    doc.save(str(f))
    doc.close()

    result = parsers.extract(f, low_text_chars_per_page=50)
    assert result.needs_ocr is True
    assert result.metadata["avg_chars_per_page"] == 0.0


def test_csv_parser(tmp_path):
    f = tmp_path / "e.csv"
    f.write_text("a,b,c\n1,2,3\n4,5,6\n")
    result = parsers.extract(f)
    assert "a,b,c" in result.text
    assert result.metadata["total_rows"] == 2
    assert result.metadata["truncated"] is False


def test_xlsx_parser(tmp_path):
    f = tmp_path / "f.xlsx"
    pd.DataFrame({"x": [1, 2], "y": [3, 4]}).to_excel(f, index=False)
    result = parsers.extract(f)
    assert "x,y" in result.text
    assert result.metadata["sheet_count"] == 1


def test_unsupported_extension_raises(tmp_path):
    f = tmp_path / "g.zip"
    f.write_bytes(b"not a real zip")
    with pytest.raises(parsers.UnsupportedFileTypeError):
        parsers.extract(f)
