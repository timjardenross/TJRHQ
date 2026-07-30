"""DOCX extraction via python-docx (USS-TJR-MSN-0205C)."""

from __future__ import annotations

import docx

from .base import ExtractionResult, ExtractionError


def extract(path) -> ExtractionResult:
    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError(f"failed to open DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return ExtractionResult(
        text=text.strip(), needs_ocr=False, page_count=None,
        metadata={"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)},
    )
